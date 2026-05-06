from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


def export_to_word(original_text, jawi_text):
    os.makedirs("output", exist_ok=True)

    doc = Document()
    doc.add_heading("Teks Bahasa Melayu ke Tulisan Jawi", level=1)

    doc.add_heading("Teks Asal Rumi", level=2)
    doc.add_paragraph(original_text)

    doc.add_heading("Teks Jawi", level=2)

    for paragraph in jawi_text.splitlines():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(paragraph)
        run.font.name = "Arial"
        run.font.size = Pt(18)

    output_path = "output/hasil_jawi.docx"
    doc.save(output_path)

    return output_path
